from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file
from github_engine import GitHubEngine
from responder_engine import ResponderEngine
from docx import Document  # Import python-docx for DOCX generation
from agent_graph import app as review_graph
import os
import uuid
import re
import markdown
from html2docx import html2docx
from bs4 import BeautifulSoup
from Markdown2docx import Markdown2docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


app = Flask(__name__)
UPLOAD_FOLDER = "upload"
REPORT_FOLDER = "reports"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["REPORT_FOLDER"] = REPORT_FOLDER
# @app.route("/", methods=["GET", "POST"])
# def index():
#     review_report = None
#     if request.method == "POST":
#         if "repo_url" in request.form and request.form["repo_url"]:
#             # Process GitHub repository
#             repo_url = request.form["repo_url"]
#             github_engine = GitHubEngine()
#             repo_name, all_chunks = github_engine.process_repository(repo_url)

#         elif "file" in request.files:
#             # Process uploaded file
#             uploaded_file = request.files["file"]
#             if uploaded_file.filename:
#                 file_path = os.path.join(app.config["UPLOAD_FOLDER"], uploaded_file.filename)
#                 uploaded_file.save(file_path)

#                 with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#                     file_content = f.read()

#                 from template_reader import TemplateReader
#                 template_reader = TemplateReader()
#                 all_chunks = template_reader.chunk_code(file_content)
#                 repo_name = uploaded_file.filename  # Use filename as repo_name

#         if all_chunks:
#             responder = ResponderEngine()
#             review_report = responder.generate_code_review(repo_name, all_chunks)

#     return render_template("intro1.html", review=review_report)
# def parse_review_sections(review_text):
#     md = MarkdownIt()
#     tokens = md.parse(review_text)

#     sections = {}
#     current_title = None
#     current_body = []

#     for token in tokens:
#         if token.type == "heading_open" and token.tag in ["h2", "h3"]:
#             # Store previous section
#             if current_title and current_body:
#                 sections[current_title.lower().replace(" ", "_")] = {
#                     "title": current_title,
#                     "body": "\n".join(current_body).strip()
#                 }
#                 current_body = []

#         elif token.type == "inline" and current_title is None:
#             current_title = token.content.strip()

#         elif token.type == "inline" and current_title:
#             current_body.append(token.content.strip())

#         elif token.type == "paragraph_open":
#             continue  # Skip — handled by `inline`

#     # Save last section
#     if current_title and current_body:
#         sections[current_title.lower().replace(" ", "_")] = {
#             "title": current_title,
#             "body": "\n".join(current_body).strip()
#         }

#     return sections

@app.route("/")
def intro():
    return render_template("intro2.html")
@app.route("/index2")
def index1():
    return render_template("index2.html")
@app.route("/review2")
def review1():
    # Load review from saved file
    report_path = os.path.join(REPORT_FOLDER,"latest_review.txt")
    if not os.path.exists(report_path):
        return render_template("review2.html", review = None)
    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()

        html = markdown.markdown(review_text)
        soup = BeautifulSoup(html, "html.parser")
        sections = {}
        current_title = None

        for tag in soup.find_all(["h2", "h3", "h4", "p", "ul", "ol", "pre", "blockquote"]):
            if tag.name.startswith("h"):
                current_title = tag.text.strip()
                sections[current_title.lower().replace(" ","_")] = {
                    "title": current_title,
                    "body": ""
                }
            elif current_title:
                sections[current_title.lower().replace(" ","_")]["body"] += str(tag)
    return render_template("review2.html", sections = sections)


# Handle Submission (POST from index1)
@app.route("/submit", methods=["POST"])
def submit():
    # all_chunks = None
    # repo_name = "Uploaded_Code"
    review_report = None

    if "repo_url" in request.form and request.form["repo_url"]:
        repo_url = request.form["repo_url"]
        github_engine = GitHubEngine()
        repo_name, all_chunks = github_engine.process_repository(repo_url)

    elif "file" in request.files:
        uploaded_file = request.files["file"]
        if uploaded_file.filename:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], uploaded_file.filename)
            uploaded_file.save(file_path)

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                file_content = f.read()

            from template_reader import TemplateReader
            template_reader = TemplateReader()
            all_chunks = template_reader.chunk_code(file_content)
            repo_name = uploaded_file.filename

    if all_chunks:
        # responder = ResponderEngine()
        result = review_graph.invoke({
            "repo": repo_name,
            "code": all_chunks
        })

        review_report = result["final_report"]

        # Save the review text
        report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(review_report)
    return jsonify({"success": True}) if review_report else jsonify({"error": "Failed to generate review"}),200



# @app.route("/download_docx", methods=["POST"])
# def download_docx():
#     data = request.get_json()
#     review_text = data.get("review", "No review available.")

#     doc = Document()
#     doc.add_heading("Code Review Report", level=1)
#     doc.add_paragraph(review_text)

#     docx_path = "review_report.docx"
#     doc.save(docx_path)

#     return send_file(docx_path, as_attachment=True)
# Download Review as DOCX
@app.route("/download_docx", methods=["POST"])
def download_docx():
    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    if not os.path.exists(report_path):
        return "Report not found", 404

    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()
    
    # project_path = os.path.join(REPORT_FOLDER, "docx_project")
    # os.makedirs(project_path, exist_ok=True)
    # docx_path = os.path.join(REPORT_FOLDER, "latest_review.docx")
    # generated_docx = os.path.join(project_path, "report.docx")
    

    # md_file_path = os.path.join(project_path, "report")
    # with open(md_file_path, 'w', encoding="utf-8") as f:
    #     f.write(review_text)
    # project_base = os.path.splitext(md_file_path)[0]
    review_html = markdown.markdown(review_text)
    soup = BeautifulSoup(review_html, "html.parser")

    doc = Document()
    doc.add_heading("Code Review Report", level=1)
    for element in soup.descendants:
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif element.name == "code":
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    # try:
    #     md2docx = Markdown2docx(project_base)
    #     md2docx.md2docx()
    #     print("Docx generated...")
    # except Exception as e:
    #     return f"Conversion failed: {str(e)}",500
    # if not os.path.exists(generated_docx):
    #     return "Docx dile not generated. check markdown formatting", 500
    # if os.path.exists(docx_path):
    #     os.remove(docx_path)
    # os.rename(generated_docx, docx_path)
    docx_path = os.path.join(REPORT_FOLDER, "latest_review.docx")
    doc.save(docx_path)

    return send_file(docx_path, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
