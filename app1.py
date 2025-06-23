from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from github_engine import GitHubEngine
from responder_engine import ResponderEngine
from parser_module import parse  # adjust if parse is defined inline
# from hf import ResponderEngine
from docx import Document
from agent_graph import agent_graph
# from agent_graph_comparision import app as comparison_graph
from conversion_graph import build_conversion_graph
from tools.bods_parser import parse_bods_script
# from conversion_graph import app as convert_graph
from bs4 import BeautifulSoup
import os
from docx.shared import Pt
import requests
import html
import openai

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

UPLOAD_FOLDER = "upload"
REPORT_FOLDER = "reports"
BODS_FOLDER="bods"
XML_FOLDER = "xml"

os.makedirs(BODS_FOLDER, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)
os.makedirs(XML_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["REPORT_FOLDER"] = REPORT_FOLDER
app.config["BODS_FOLDER"] = BODS_FOLDER
app.config["XML_FOLDER"] = XML_FOLDER

@app.route("/submit", methods=["POST"])
def submit():
    review_report = None
    comparison_report = None
    used_agents = []

    azure_key = request.form.get('azureKey')
    azure_endpoint = request.form.get('azureEndpoint')
    azure_deployment = request.form.get('azureDeployment')
    azure_api_version = request.form.get('azureApiVersion')
 
    if azure_key and azure_endpoint and azure_deployment and azure_api_version:
        openai.api_type = "azure"
        openai.api_key = azure_key
        openai.api_base = azure_endpoint
        openai.api_version = azure_api_version

    # Get files
    uploaded_files = request.files.getlist("file")
    groundtruth_file = request.files.get("ground_truth")
    file_uploaded = uploaded_files and uploaded_files[0].filename
    groundtruth_uploaded = groundtruth_file and groundtruth_file.filename

    if "repo_url" in request.form and request.form["repo_url"]:
        repo_url = request.form["repo_url"]
        github_engine = GitHubEngine()
        repo_name, all_chunks = github_engine.process_repository(repo_url)
        used_agents += ['qualityAnalyzer', 'bugDetector']
            
        responder = ResponderEngine()
        review_app = agent_graph(responder)
        result = review_app.invoke({
            "repo": repo_name,
            "code": all_chunks
        })

        review_report = result["final_report"]

    elif "file" in request.files:
        print("file uploaded")
        uploaded_file = request.files["file"]
        if uploaded_file.filename:
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], uploaded_file.filename)
            uploaded_file.save(file_path)

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                file_content = f.read()

            from chunking_engine import Chunking
            chunking_engine = Chunking()
            all_chunks = chunking_engine.chunk_code(file_content)
            repo_name = uploaded_file.filename
            used_agents += ['improvementSuggestor', 'unitTestSuggestor']
            responder = ResponderEngine()
            review_app = agent_graph(responder)
            result = review_app.invoke({
                "repo": repo_name,
                "code": all_chunks
            })
            review_report = result["final_report"]
    else:
        return jsonify({"error": "No valid input provided"}),400

    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(review_report)
    return jsonify({
        "success": True,
        "review_report": review_report,
        "used_agents": used_agents
    }), 200


@app.route("/migrate", methods=["POST"])
def migrate():
    comparison_report = None
    used_agents = []
 
    azure_key = request.form.get('azureKey')
    azure_endpoint = request.form.get('azureEndpoint')
    azure_deployment = request.form.get('azureDeployment')
    azure_api_version = request.form.get('azureApiVersion')
 
    if azure_key and azure_endpoint and azure_deployment and azure_api_version:
        openai.api_type = "azure"
        openai.api_key = azure_key
        openai.api_base = azure_endpoint
        openai.api_version = azure_api_version
 
    file_uploaded = "file" in request.files and request.files["file"].filename
    groundtruth_uploaded = "ground_truth" in request.files and request.files["ground_truth"].filename
 
    if file_uploaded and groundtruth_uploaded:
        groundtruth_file = request.files["ground_truth"]
        converted_file = request.files["file"]
 
        groundtruth_path = os.path.join(app.config["UPLOAD_FOLDER"], groundtruth_file.filename)
        converted_path = os.path.join(app.config["UPLOAD_FOLDER"], converted_file.filename)
 
        groundtruth_file.save(groundtruth_path)
        converted_file.save(converted_path)
 
        with open(groundtruth_path, "r", encoding="utf-8", errors="ignore") as f:
            original_code = f.read()
 
        with open(converted_path, "r", encoding="utf-8", errors="ignore") as f:
            converted_code = f.read()
        from chunking_engine import Chunking
        chunking_engine = Chunking()
 
        original_chunks = chunking_engine.chunk_code(original_code)
        converted_chunks = chunking_engine.chunk_code(converted_code)

        # Initialize and call the responder engine
        from responder_engine import ResponderEngine
        responder = ResponderEngine()
        comparison_report = responder.run_functionality_comparison(original_chunks, converted_chunks)
 
        report_path = os.path.join(REPORT_FOLDER, "latest_comparison.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(comparison_report)
 
        used_agents.append('responderEngine')
 
    return jsonify({
        "success": True,
        "comparison_report": comparison_report,
        "used_agents": used_agents
    }), 200

@app.route("/convert", methods=["POST"])
def convert():
    converted_code = None
    used_agents = []

    BODS_uploaded = "bods_file" in request.files and request.files["bods_file"].filename
    XML_uploaded = "xml_file" in request.files and request.files["xml_file"].filename
    if BODS_uploaded and XML_uploaded:
        bods_file = request.files["bods_file"]
        bods_path = os.path.join(app.config["BODS_FOLDER"], bods_file.filename)
        bods_file.save(bods_path)

        xml_file = request.files["xml_file"]
        xml_path = os.path.join(app.config["XML_FOLDER"], xml_file.filename)
        xml_file.save(xml_path)

        target = request.form["target_language"]
        print("Target: ", target)
 
        with open(bods_path, "r", encoding="utf-8", errors="ignore") as f:
            original_code = f.read()

        with open(xml_path, "r", encoding= "utf-8", errors= "ignore") as f:
            xml_code = f.read()  
        graph = build_conversion_graph()
        # parsed = parse_bods_script(original_code)
        # print(parsed)
        result = graph.invoke({
            "bods_script": original_code,
            "target" : target,
            "xml_code": xml_code
            })
        print("Extraction:",result.get("extraction_result",""))
        print("Job:",result.get("job_result",""))
        output = (
            result.get("extraction_result", "")+ "\n\n" +
            result.get("job_result", "")
        )
        converted_code = parse(output)

        # from responder_engine import ResponderEngine
        # responder = ResponderEngine()
        # converted_code = responder.converter(original_chunks)

 
        # Create 'converted' folder if it doesn't exist
        converted_folder = os.path.join(os.getcwd(), "converted")
        os.makedirs(converted_folder, exist_ok=True)
 
        # Determine output filename
        base_filename = os.path.splitext(bods_file.filename)[0]
        if target == "SQL":
            converted_filename = f"{base_filename}_talend.sql"
        elif target == "Pyspark":
            converted_filename = f"{base_filename}_talend.py"
        elif target == "Scala":
            converted_filename = f"{base_filename}_talend.scala"
        else:
            converted_filename = "SAP BODS to talend.java"

        converted_path = os.path.join(converted_folder, converted_filename)
 
        # Save the converted code to the file
        with open(converted_path, "w", encoding="utf-8") as out_file:
            out_file.write(converted_code)
 
        # return send_from_directory(
        #     directory= converted_folder,
        #     path= converted_filename,
        #     as_attachment = True
        # )
 
        return jsonify({
            "message": "Conversion successful",
            "converted_file": converted_filename,
            "converted_code": converted_code
        })
    else:
        bods_file = request.files["bods_file"]
        bods_path = os.path.join(app.config["BODS_FOLDER"], bods_file.filename)
        bods_file.save(bods_path)

        target = request.form["target_language"]
        print("Target: ", target)
 
        with open(bods_path, "r", encoding="utf-8", errors="ignore") as f:
            original_code = f.read()
        graph = build_conversion_graph()
        result = graph.invoke({
            "bods_script": original_code,
            "target" : target
            })
        # converted_code = parse(result["extraction_result"])
        converted_code = (
            result.get("extraction_result", "")+ "\n\n" +
            result.get("job_result", "")
        )
        converted_folder = os.path.join(os.getcwd(), "converted")
        os.makedirs(converted_folder, exist_ok=True)
 
        # Determine output filename
        base_filename = os.path.splitext(bods_file.filename)[0]
        if target == "SQL":
            converted_filename = f"{base_filename}_talend.sql"
        elif target == "Pyspark":
            converted_filename = f"{base_filename}_talend.py"
        elif target == "Scala":
            converted_filename = f"{base_filename}_talend.scala"
        else:
            converted_filename = "SAP BODS to talend.java"

        converted_path = os.path.join(converted_folder, converted_filename)
 
        # Save the converted code to the file
        with open(converted_path, "w", encoding="utf-8") as out_file:
            out_file.write(converted_code)
        return jsonify({
            "message": "Conversion successful",
            "converted_file": converted_filename,
            "converted_code": converted_code
        })
    
@app.route('/download/<filename>', methods = ['GET'])
def download_file(filename):
    converted_folder = os.path.join(os.getcwd(), 'converted')
    return send_from_directory(converted_folder, filename, as_attachment=True)

@app.route("/analyze_selected", methods=["POST"])
def analyze_selected():
    data = request.get_json()
    file_urls = data.get("file_urls", [])
 
    if not file_urls:
        return jsonify({"success": False, "error": "No files provided."})
 
    file_contents_combined = ""
    for url in file_urls:
        response = requests.get(url)
        if response.ok:
            file_contents_combined += response.text + "\n\n"
        else:
            return jsonify({"success": False, "error": f"Failed to fetch {url}."})
 
    from chunking_engine import Chunking
    chunking_engine = Chunking()
    all_chunks = chunking_engine.chunk_code(file_contents_combined)    
    responder = ResponderEngine()
    review_app = agent_graph(responder)
 
    if all_chunks:
        result = review_app.invoke({
            "repo": "Selected_Files",
            "code": all_chunks
        })
 
        review_report = result["final_report"]
 
        # Save review
        report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(review_report)
 
        return jsonify({"success": True})
 
    return jsonify({"success": False, "error": "Failed to process selected files."})

@app.route("/review2", methods=["GET"])
def review2():
    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    if not os.path.exists(report_path):
        return jsonify({"error": "No report found."}), 404

    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()

    soup = BeautifulSoup(review_text, "html.parser")
    sections = []
    current_section = None

    for tag in soup.find_all(["h3", "p", "ul", "ol", "pre", "blockquote"]):
        if tag.name == "h3":
            if current_section:
                sections.append(current_section)
            current_section = {
                "title": tag.get_text(strip=True),
                "body": ""
            }
        elif current_section:
            if tag.name == "pre":
                code_tag = tag.find("code")
                if code_tag:
                    raw_code = code_tag.decode_contents()
                    escaped_code = html.escape(raw_code)
                    lang_class = code_tag.get("class", [""])[0]
                    current_section["body"] += f"<pre><code class='{lang_class}'>{escaped_code}</code></pre>\n"
                else:
                    raw_code = tag.get_text()
                    current_section["body"] += f"<pre><code>{html.escape(raw_code)}</code></pre>\n"
            else:
                current_section["body"] += str(tag)

    if current_section:
        sections.append(current_section)

    return jsonify(sections)
    # comparison_path = os.path.join(REPORT_FOLDER, "latest_comparison.txt")

@app.route('/compare', methods=["GET"])
def compare_report():
    # Load the HTML report
    comparison_path = os.path.join(REPORT_FOLDER, "latest_comparison.txt")
    if not os.path.exists(comparison_path):
        return jsonify({"error": "No report found."}), 404
    with open(comparison_path, 'r', encoding='utf-8') as f:
        html = f.read()

    soup = BeautifulSoup(html, 'html.parser')

    functional_table_element = soup.find('table')
    overall_conversion_element = functional_table_element.find_next_sibling('p')

    functional_table = str(functional_table_element)
    overall_conversion = str(overall_conversion_element)

    # Helper function to collect all elements between current h3 and the next h3
    def collect_section_content(header_text):
        section = ''
        h3 = soup.find('h3', string=header_text)
        if h3:
            next_tag = h3.find_next_sibling()
            while next_tag and next_tag.name != 'h3':
                section += str(next_tag)
                next_tag = next_tag.find_next_sibling()
        return section

    code_quality = collect_section_content("Code Quality Assessment")
    suggestions = collect_section_content("Enhancements / Suggestions")

    # Revised Code
    revised_code = ''.join(str(pre) for pre in soup.select('h3:contains("Revised Code") ~ pre'))

    return jsonify({
        "functional_table": functional_table + overall_conversion,
        "code_quality": code_quality,
        "suggestions": suggestions,
        "revised_code": revised_code
    })
    # code_quality = ''.join(str(ul) for ul in soup.select('h3:contains("Code Quality Assessment") ~ ul'))
    # suggestions = ''.join(str(ul) for ul in soup.select('h3:contains("Enhancements / Suggestions") ~ ul'))
    # revised_code = ''.join(str(pre) for pre in soup.select('h3:contains("Revised Code") ~ pre'))

    # return jsonify({
    #     "functional_table": functional_table + overall_conversion,
    #     "code_quality": code_quality,
    #     "suggestions": suggestions,
    #     "revised_code": revised_code
    # })



@app.route("/download_docx", methods=["GET"])
def download_docx():
    report_path = os.path.join(REPORT_FOLDER, "latest_review.txt")
    if not os.path.exists(report_path):
        return jsonify({"error": "Report not found"}), 404

    with open(report_path, "r", encoding="utf-8") as f:
        review_text = f.read()
    
    soup = BeautifulSoup(review_text, "html.parser")
    doc = Document()
    doc.add_heading("Code Review Report", level=1)

    for element in soup.find_all(True):
        tag_name = element.name.lower()
        text = element.get_text(strip=True)

        if not text:
            continue  # skip empty elements

        if tag_name == "h1":
            doc.add_heading(text, level=1)
        elif tag_name == "h2":
            doc.add_heading(text, level=2)
        elif tag_name == "h3":
            doc.add_heading(text, level=3)
        elif tag_name == "h4":
            doc.add_heading(text, level=4)
        elif tag_name == "p":
            doc.add_paragraph(text)
        elif tag_name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif tag_name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif tag_name == "pre":
            code_text = element.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
        elif tag_name == "code":
            run = doc.add_paragraph().add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    docx_path = os.path.join(REPORT_FOLDER, "CodeSphere_Review_Report.docx")
    doc.save(docx_path)

    return send_file(
        docx_path,
        as_attachment=True,
        download_name="CodeSphere_Review_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/download_comparison_report", methods=["GET"])
def compare_download():
    compare_path = os.path.join(REPORT_FOLDER, "latest_comparison.txt")
    if not os.path.exists(compare_path):
        return jsonify({"error": "Report not found"}), 404

    with open(compare_path, "r", encoding="utf-8") as f:
        compare_text = f.read()
    soup = BeautifulSoup(compare_text, "html.parser")
    doc = Document()
    doc.add_heading("Code Review Report", level=1)

    for element in soup.find_all(True):
        tag_name = element.name.lower()
        text = element.get_text(strip=True)

        if not text:
            continue  # skip empty elements

        if tag_name == "h1":
            doc.add_heading(text, level=1)
        elif tag_name == "h2":
            doc.add_heading(text, level=2)
        elif tag_name == "h3":
            doc.add_heading(text, level=3)
        elif tag_name == "h4":
            doc.add_heading(text, level=4)
        elif tag_name == "p":
            doc.add_paragraph(text)
        elif tag_name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif tag_name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif tag_name == "pre":
            code_text = element.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
        elif tag_name == "code":
            run = doc.add_paragraph().add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif tag_name == "table":
            # Create a new table
            rows = element.find_all("tr")
            if rows:
                # Determine the number of columns from the first row
                num_columns = len(rows[0].find_all(['th', 'td']))
                table = doc.add_table(rows=len(rows), cols=num_columns)
                table.style = 'Table Grid'  # Optional: apply a style

                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text(strip=True)


    docx_path = os.path.join(REPORT_FOLDER, "CodeSphere_Comparison_Report.docx")
    doc.save(docx_path)

    return send_file(
        docx_path,
        as_attachment=True,
        download_name="CodeSphere_Comparison_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/")
def health_check():
    return jsonify({"message": "Flask backend is running."})


if __name__ == "__main__":
    app.run(debug=True, use_reloader= False)